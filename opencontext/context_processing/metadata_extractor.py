#!/usr/bin/env python
# -*- coding: utf-8 -*-

# Copyright (c) 2025 Beijing Volcano Engine Technology Co., Ltd.
# SPDX-License-Identifier: Apache-2.0

"""
Metadata Extraction Service
Extracts comprehensive metadata from various file types including:
- File properties (size, dates, permissions)
- Author and creator information
- Document properties (title, subject, keywords)
- Technical metadata (format, encoding, dimensions)
"""

import os
import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
import mimetypes

from opencontext.utils.logging_utils import get_logger

logger = get_logger(__name__)


class MetadataExtractor:
    """
    Service for extracting metadata from various file types
    """

    def __init__(self):
        self._extractors = {
            "pdf": self._extract_pdf_metadata,
            "docx": self._extract_docx_metadata,
            "xlsx": self._extract_xlsx_metadata,
            "image": self._extract_image_metadata,
            "code": self._extract_code_metadata,
        }

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract comprehensive metadata from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing extracted metadata
        """
        try:
            path = Path(file_path)
            
            if not path.exists():
                logger.warning(f"File not found: {file_path}")
                return {}
            
            # Basic file metadata
            metadata = self._extract_basic_metadata(path)
            
            # File type-specific metadata
            file_ext = path.suffix.lower()
            file_type = self._detect_file_type(file_ext)
            
            if file_type in self._extractors:
                try:
                    specific_metadata = self._extractors[file_type](path)
                    metadata.update(specific_metadata)
                except Exception as e:
                    logger.warning(f"Error extracting {file_type} metadata from {file_path}: {e}")
            
            return metadata
            
        except Exception as e:
            logger.exception(f"Error extracting metadata from {file_path}: {e}")
            return {}

    def _extract_basic_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract basic file system metadata"""
        stat = path.stat()
        
        metadata = {
            "file_name": path.name,
            "file_path": str(path.absolute()),
            "file_size": stat.st_size,
            "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
            "created_time": datetime.datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_time": datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "accessed_time": datetime.datetime.fromtimestamp(stat.st_atime).isoformat(),
            "file_extension": path.suffix.lower(),
            "file_stem": path.stem,
        }
        
        # MIME type
        mime_type, _ = mimetypes.guess_type(str(path))
        if mime_type:
            metadata["mime_type"] = mime_type
        
        # File permissions (Unix-like systems)
        try:
            metadata["permissions"] = oct(stat.st_mode)[-3:]
        except:
            pass
        
        return metadata

    def _detect_file_type(self, extension: str) -> Optional[str]:
        """Detect general file type category"""
        if extension == ".pdf":
            return "pdf"
        elif extension in [".docx", ".doc"]:
            return "docx"
        elif extension in [".xlsx", ".xls"]:
            return "xlsx"
        elif extension in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"]:
            return "image"
        elif extension in [".py", ".js", ".java", ".cpp", ".c", ".go", ".rs"]:
            return "code"
        return None

    def _extract_pdf_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files"""
        try:
            from pypdf import PdfReader
            
            metadata = {}
            reader = PdfReader(str(path))
            
            # Document info
            if reader.metadata:
                for key, value in reader.metadata.items():
                    clean_key = key.lstrip('/')
                    metadata[f"pdf_{clean_key.lower()}"] = str(value) if value else None
            
            # Page count
            metadata["pdf_page_count"] = len(reader.pages)
            
            # Extract text from first page for preview
            if len(reader.pages) > 0:
                first_page_text = reader.pages[0].extract_text()
                metadata["pdf_first_page_preview"] = first_page_text[:500] if first_page_text else None
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting PDF metadata: {e}")
            return {}

    def _extract_docx_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX files"""
        try:
            from docx import Document
            
            metadata = {}
            doc = Document(str(path))
            
            # Core properties
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["docx_title"] = core_props.title
            if core_props.author:
                metadata["docx_author"] = core_props.author
            if core_props.subject:
                metadata["docx_subject"] = core_props.subject
            if core_props.keywords:
                metadata["docx_keywords"] = core_props.keywords
            if core_props.created:
                metadata["docx_created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["docx_modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata["docx_last_modified_by"] = core_props.last_modified_by
            
            # Document statistics
            metadata["docx_paragraph_count"] = len(doc.paragraphs)
            metadata["docx_table_count"] = len(doc.tables)
            
            # Extract text preview
            text_preview = "\n".join([p.text for p in doc.paragraphs[:5] if p.text])
            metadata["docx_text_preview"] = text_preview[:500] if text_preview else None
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata: {e}")
            return {}

    def _extract_xlsx_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract metadata from Excel files"""
        try:
            import openpyxl
            
            metadata = {}
            workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            
            # Workbook properties
            props = workbook.properties
            if props.title:
                metadata["xlsx_title"] = props.title
            if props.creator:
                metadata["xlsx_creator"] = props.creator
            if props.subject:
                metadata["xlsx_subject"] = props.subject
            if props.keywords:
                metadata["xlsx_keywords"] = props.keywords
            if props.created:
                metadata["xlsx_created"] = props.created.isoformat()
            if props.modified:
                metadata["xlsx_modified"] = props.modified.isoformat()
            if props.lastModifiedBy:
                metadata["xlsx_last_modified_by"] = props.lastModifiedBy
            
            # Sheet information
            metadata["xlsx_sheet_count"] = len(workbook.sheetnames)
            metadata["xlsx_sheet_names"] = workbook.sheetnames
            
            # Statistics from first sheet
            if workbook.sheetnames:
                first_sheet = workbook[workbook.sheetnames[0]]
                metadata["xlsx_first_sheet_rows"] = first_sheet.max_row
                metadata["xlsx_first_sheet_cols"] = first_sheet.max_column
            
            workbook.close()
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting XLSX metadata: {e}")
            return {}

    def _extract_image_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract metadata from image files"""
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            
            metadata = {}
            
            with Image.open(str(path)) as img:
                # Basic image info
                metadata["image_format"] = img.format
                metadata["image_mode"] = img.mode
                metadata["image_width"] = img.width
                metadata["image_height"] = img.height
                metadata["image_size"] = f"{img.width}x{img.height}"
                
                # EXIF data
                exif = img.getexif()
                if exif:
                    exif_data = {}
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        # Only include string/number values
                        if isinstance(value, (str, int, float)):
                            exif_data[f"exif_{tag}"] = value
                    metadata.update(exif_data)
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting image metadata: {e}")
            return {}

    def _extract_code_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract metadata from code files"""
        try:
            metadata = {}
            
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                lines = content.split('\n')
            
            # Line statistics
            metadata["code_total_lines"] = len(lines)
            metadata["code_non_empty_lines"] = sum(1 for line in lines if line.strip())
            metadata["code_comment_lines"] = sum(1 for line in lines if line.strip().startswith(('#', '//', '/*', '*')))
            
            # File encoding detection
            import chardet
            with open(path, 'rb') as f:
                raw = f.read()
                result = chardet.detect(raw)
                metadata["code_encoding"] = result['encoding']
                metadata["code_encoding_confidence"] = result['confidence']
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting code metadata: {e}")
            return {}


# Global instance
_metadata_extractor = None


def get_metadata_extractor() -> MetadataExtractor:
    """Get global metadata extractor instance"""
    global _metadata_extractor
    if _metadata_extractor is None:
        _metadata_extractor = MetadataExtractor()
    return _metadata_extractor
